import json
import os
import threading
from datetime import date, datetime, timedelta
from io import BytesIO

import streamlit as st

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
ENTRIES_FILE = os.path.join(DATA_DIR, "entries.json")
NOTES_FILE = os.path.join(DATA_DIR, "manager_notes.json")
EXPORTS_DIR = os.path.join(DATA_DIR, "exports")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(EXPORTS_DIR, exist_ok=True)
LOCK = threading.Lock()

SEED = {
    "2026-08-03": {"date": "2026-08-03", "tasks": ["Reviewed, edited and sent 13 T18 documents to Alok", "Worked on Conval and Excellis Audit Report spreadsheet for Pooja"], "goals": "- Work on more T18 documents", "hours": 7, "status": "on-track", "notes": "", "learning": "", "idea": "", "reflection": "", "private": True},
    "2026-08-04": {"date": "2026-08-04", "tasks": ["Reviewed, edited and sent all remaining T18 documents", "Finished some Litmos courses", "Meeting with Alok, Pooja, Deepak on Partner Scorecard Dashboard"], "goals": "Set clearer goal for next day", "hours": 6, "status": "on-track", "notes": "Which Litmos course, how much time per activity — need to specify daily goals better", "learning": "", "idea": "", "reflection": "", "private": True},
    "2026-08-05": {"date": "2026-08-05", "tasks": ["Litmos courses completed: FOSCM, QMS Privacy & Data Protection, NS Overview for New Hires, QMS CS01 Technical Support, Preventing Harassment (Global Workplace), QMS Data Integrity Policy"], "goals": "- Complete QMS PS01 course\n- Complete more Litmos courses\n- Prepare and attend SPE Monthly Meet\n- Attend Use Org (Services Training) Meet", "hours": 7, "status": "on-track", "notes": "", "learning": "", "idea": "", "reflection": "", "private": True},
    "2026-08-06": {"date": "2026-08-06", "tasks": ["Meetings attended: SPE Monthly Meeting, Use Org (Services Training) Meet.", "Litmos completed: QMS PS01, SVC Use Org Training, Intro to GS1 Standards, OPUS No Code Fundamentals, MINT External Manufacturing Configuration"], "goals": "- Complete all remaining Litmos courses before 8 Aug deadline\n- Weekly catchup with manager", "hours": 7, "status": "under", "notes": "Litmos deadline is tight — running behind on hours logged mid-week", "learning": "", "idea": "", "reflection": "", "private": True},
}


def normalize_tasks(value):
    if isinstance(value, list):
        return [str(x).strip().lstrip("-•* ").strip() for x in value if str(x).strip()]
    text = str(value or "").strip()
    if not text:
        return []
    return [line.strip().lstrip("-•* ").strip() for line in text.splitlines() if line.strip()]


def normalize_entries(entries):
    changed = False
    for iso, entry in entries.items():
        tasks = normalize_tasks(entry.get("tasks"))
        if entry.get("tasks") != tasks:
            entry["tasks"] = tasks
            changed = True
    return entries, changed


def load_json(path, default):
    with LOCK:
        if not os.path.exists(path):
            with open(path, "w", encoding="utf-8") as f:
                json.dump(default, f, indent=2, ensure_ascii=False)
            return default
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return default


def save_json(path, value):
    with LOCK:
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(value, f, indent=2, ensure_ascii=False)
        os.replace(tmp, path)


def load_entries():
    if not os.path.exists(ENTRIES_FILE):
        save_json(ENTRIES_FILE, SEED)
        return dict(SEED)
    entries = load_json(ENTRIES_FILE, {})
    entries, changed = normalize_entries(entries)
    if changed:
        save_json(ENTRIES_FILE, entries)
    return entries


def load_notes():
    return load_json(NOTES_FILE, [])


def save_entry(entry):
    entries = load_entries()
    entry = dict(entry)
    entry["tasks"] = normalize_tasks(entry.get("tasks"))
    entries[entry["date"]] = entry
    save_json(ENTRIES_FILE, entries)


def add_manager_note(day_iso, author, text):
    notes = load_notes()
    notes.append({
        "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
        "date": day_iso,
        "text": text.strip(),
        "author": author.strip() or "Manager",
        "ts": datetime.now().strftime("%Y-%m-%d %H:%M"),
    })
    save_json(NOTES_FILE, notes)


def notes_for_day(notes, day_iso):
    return [n for n in notes if n.get("date") == day_iso]


def day_name(iso):
    return datetime.strptime(iso, "%Y-%m-%d").strftime("%A")


def fmt_date(iso):
    return datetime.strptime(iso, "%Y-%m-%d").strftime("%d %b %Y")


def week_start(iso):
    d = datetime.strptime(iso, "%Y-%m-%d").date()
    return (d - timedelta(days=d.weekday())).isoformat()


def tasks_text(tasks):
    return "\n".join(f"- {task}" for task in normalize_tasks(tasks))


def inject_css():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,600;9..144,700;9..144,800&family=IBM+Plex+Sans:wght@400;500;600;700&family=IBM+Plex+Mono:wght@400;500;600;700&display=swap');
        :root{--purple:#1B4F91;--purple-dark:#0A2A52;--teal:#2E8FD6;--pink:#0E7C86;--line:#D8EEF4;}
        html,body,[class*="css"]{font-family:'IBM Plex Sans',sans-serif;}
        .stApp{background:radial-gradient(circle at 8% 0%,rgba(236,72,153,.12),transparent 40%),radial-gradient(circle at 100% 10%,rgba(46,143,214,.12),transparent 45%),#F6F3FF;}
        h1,h2,h3{font-family:'Fraunces',serif !important;}
        .hero{background:linear-gradient(120deg,#0A2A52,#1B4F91 45%,#2E8FD6 85%,#0E7C86);padding:28px 32px;border-radius:0 0 22px 22px;color:white;margin:-1rem -1rem 1.5rem -1rem;box-shadow:0 8px 32px rgba(10,42,82,.25);}
        .hero-eyebrow{font:700 12px 'IBM Plex Mono',monospace;letter-spacing:.12em;text-transform:uppercase;color:#DCEEFF;}
        .hero h1{font-size:44px;margin:4px 0;color:white !important;}
        .hero p{margin:0;color:#EAF7FA;}
        .task-card{background:white;border:1px solid #D8EEF4;border-radius:14px;padding:10px 12px;margin-bottom:8px;box-shadow:0 2px 8px rgba(27,79,145,.06);}
        .task-number{font:700 12px 'IBM Plex Mono',monospace;color:#1B4F91;padding-top:9px;}
        .manager-banner{background:#DFEEFB;border:1px solid #B7E7EF;border-radius:12px;padding:14px 16px;color:#0A5A6B;}
        .task-list{margin:0;padding-left:20px;}
        .task-list li{margin-bottom:6px;line-height:1.45;}
        .small-muted{color:#6b7280;font-size:.85rem;}

        /* --- Manager view redesign --- */
        .week-banner{display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:10px;background:linear-gradient(100deg,#0A2A52,#1B4F91 60%,#0E7C86);border-radius:16px;padding:16px 22px;margin:22px 0 14px 0;color:#fff;box-shadow:0 6px 20px rgba(10,42,82,.18);}
        .week-banner .wtitle{font-family:'Fraunces',serif;font-size:20px;font-weight:700;}
        .week-banner .wsub{font:600 12px 'IBM Plex Mono',monospace;letter-spacing:.06em;color:#DCEEFF;text-transform:uppercase;margin-top:2px;}
        .week-chip{background:rgba(255,255,255,.14);border:1px solid rgba(255,255,255,.28);border-radius:999px;padding:6px 14px;font:700 13px 'IBM Plex Mono',monospace;}
        .day-card{background:#fff;border-radius:14px;padding:16px 18px;margin-bottom:14px;box-shadow:0 2px 10px rgba(10,42,82,.07);border-left:6px solid #9CA3AF;}
        .day-card.status-on{border-left-color:#1FA971;}
        .day-card.status-under{border-left-color:#E08A2E;}
        .day-card.status-blocked{border-left-color:#D6455C;}
        .day-head{display:flex;align-items:baseline;justify-content:space-between;flex-wrap:wrap;gap:8px;margin-bottom:6px;}
        .day-date{font-family:'Fraunces',serif;font-weight:700;font-size:17px;color:#0A2A52;}
        .day-name{color:#6b7280;font-size:.82rem;margin-left:6px;}
        .badge{display:inline-block;padding:4px 12px;border-radius:999px;font:700 12px 'IBM Plex Mono',monospace;letter-spacing:.02em;}
        .badge-on{background:#E4F8EE;color:#0E7A47;}
        .badge-under{background:#FDF0DF;color:#B4620C;}
        .badge-blocked{background:#FCE7EA;color:#B22C42;}
        .hours-chip{background:#F1F0FF;color:#4B3F91;font:700 12px 'IBM Plex Mono',monospace;padding:4px 10px;border-radius:8px;margin-left:8px;}
        .section-label{font:700 11px 'IBM Plex Mono',monospace;text-transform:uppercase;letter-spacing:.08em;color:#8B93A7;margin:10px 0 4px 0;}
        .note-bubble{background:#FFF7E6;border:1px solid #F3DFAE;border-radius:10px;padding:10px 14px;margin-top:8px;}
        .note-bubble .note-meta{font:700 11px 'IBM Plex Mono',monospace;color:#946B0C;text-transform:uppercase;letter-spacing:.05em;margin-bottom:2px;}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header(manager=False):
    label = "MANAGER VIEW" if manager else "YOUR DASHBOARD"
    subtitle = "Read-only overview of Saniya's daily progress" if manager else "Daily ops tracking & working journal — Saniya"
    st.markdown(f'<div class="hero"><div class="hero-eyebrow">Intern · TraceLink · {label}</div><div class="hero h1"><h1>Field Log</h1></div><p>{subtitle}</p></div>', unsafe_allow_html=True)


def render_task_list(tasks):
    tasks = normalize_tasks(tasks)
    if not tasks:
        st.caption("No tasks logged.")
        return
    st.markdown('<ul class="task-list">' + ''.join(f'<li>{task.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")}</li>' for task in tasks) + '</ul>', unsafe_allow_html=True)


def render_stats(entries):
    today = date.today()
    ws = (today - timedelta(days=today.weekday())).isoformat()
    week_entries = [e for d,e in entries.items() if week_start(d) == ws]
    hours = sum(float(e.get("hours") or 0) for e in week_entries)
    on = sum(e.get("status") == "on-track" for e in week_entries)
    under = sum(e.get("status") == "under" for e in week_entries)
    blocked = sum(e.get("status") == "blocked" for e in week_entries)
    cols = st.columns(6)
    vals = [(len(week_entries),"Entries this week"),(hours,"Hours this week"),(on,"On track"),(under,"Under target"),(blocked,"Blocked"),(len(entries),"Total days")]
    for col,(v,lbl) in zip(cols,vals):
        col.metric(lbl, v)


def render_entry(entries):
    st.header("Log a day")
    st.caption("Add each completed task as its own bullet. Use + Add task when you have more work to record.")
    dates = sorted(entries.keys())
    default = date.today()
    if dates and default.isoformat() not in entries:
        default = date.today()
    selected = st.date_input("Date", value=default, key="entry_date")
    iso = selected.isoformat()
    existing = entries.get(iso, {})
    st.caption(f"{day_name(iso)} · {fmt_date(iso)}")

    existing_tasks = normalize_tasks(existing.get("tasks"))
    if "task_rows" not in st.session_state or st.session_state.get("task_date") != iso:
        st.session_state.task_rows = existing_tasks[:] if existing_tasks else [""]
        st.session_state.task_date = iso

    st.subheader("Ops — visible to manager")
    st.markdown("**Tasks completed**")
    st.caption("One task per line, with separate inputs for cleaner bullets in the Manager View.")
    remove_idx = None
    for i in range(len(st.session_state.task_rows)):
        c1, c2, c3 = st.columns([0.06, 0.84, 0.10])
        with c1:
            st.markdown(f'<div class="task-number">{i+1:02}</div>', unsafe_allow_html=True)
        with c2:
            st.session_state.task_rows[i] = st.text_input("", value=st.session_state.task_rows[i], key=f"task_{iso}_{i}", placeholder="e.g. Reviewed and sent T18 documents", label_visibility="collapsed")
        with c3:
            if len(st.session_state.task_rows) > 1 and st.button("✕", key=f"remove_{iso}_{i}"):
                remove_idx = i
    if remove_idx is not None:
        st.session_state.task_rows.pop(remove_idx)
        st.rerun()
    if st.button("＋ Add task", key=f"add_task_{iso}"):
        st.session_state.task_rows.append("")
        st.rerun()

    goals = st.text_area("Goals for tomorrow", value=existing.get("goals", ""), placeholder="- Finish remaining Litmos courses\n- Prep notes for weekly catchup")
    c1,c2 = st.columns(2)
    with c1:
        hours = st.number_input("Hours spent", min_value=0.0, max_value=16.0, step=0.5, value=float(existing.get("hours") or 0))
    with c2:
        status_options = ["on-track","under","blocked"]
        current_status = existing.get("status", "on-track") if existing.get("status") in status_options else "on-track"
        status = st.selectbox("Status", status_options, index=status_options.index(current_status), format_func=lambda x: {"on-track":"🟢 On track","under":"🟠 Under target","blocked":"🔴 Blocked"}[x])
    notes = st.text_area("Notes / blockers (optional)", value=existing.get("notes", ""), placeholder="Anything slowing you down or worth flagging")

    st.subheader("Journal — your working reflection")
    st.caption("Everything below is optional — save with just the Ops section if that's all you have time for.")
    learning = st.text_area("Key learning today (optional)", value=existing.get("learning", ""))
    idea = st.text_area("Idea / spark (optional)", value=existing.get("idea", ""))
    reflection = st.text_area("Thoughts / reflection (optional)", value=existing.get("reflection", ""))
    private = st.checkbox("Keep this reflection private (excluded from Manager View)", value=existing.get("private", True))

    if st.button("Save entry", type="primary", use_container_width=True):
        clean_tasks = [x.strip() for x in st.session_state.task_rows if x.strip()]
        save_entry({"date":iso,"tasks":clean_tasks,"goals":goals,"hours":hours,"status":status,"notes":notes,"learning":learning,"idea":idea,"reflection":reflection,"private":private})
        st.session_state["just_saved"] = iso
        st.rerun()

    if st.session_state.get("just_saved") == iso:
        st.success(f"✅ Entry for {fmt_date(iso)} saved — {datetime.now().strftime('%I:%M %p')}")
        del st.session_state["just_saved"]


def render_ops_log(entries):
    st.header("Ops Log")
    if not entries:
        st.info("No entries yet.")
        return
    for iso in sorted(entries.keys(), reverse=True):
        e = entries[iso]
        with st.expander(f"{fmt_date(iso)} · {day_name(iso)} · {e.get('hours',0)}h · {e.get('status','on-track').replace('-',' ').title()}", expanded=(iso == date.today().isoformat())):
            render_task_list(e.get("tasks"))
            st.write(f"**Goals:** {e.get('goals') or '—'}")
            st.write(f"**Notes / blockers:** {e.get('notes') or '—'}")


def render_journal(entries):
    st.header("Journal")
    rows = [(d,e) for d,e in sorted(entries.items(), reverse=True) if e.get("learning") or e.get("idea") or e.get("reflection")]
    if not rows:
        st.info("No journal reflections yet.")
        return
    for iso,e in rows:
        with st.expander(f"{fmt_date(iso)} · {day_name(iso)}"):
            if e.get("learning"): st.write("**Key learning**", e["learning"])
            if e.get("idea"): st.write("**Idea / spark**", e["idea"])
            if e.get("reflection"): st.write("**Reflection**", e["reflection"])
            if e.get("private", True): st.caption("Reflection marked private in the Manager View.")


STATUS_META = {
    "on-track": {"cls": "status-on", "badge": "badge-on", "label": "🟢 On track"},
    "under": {"cls": "status-under", "badge": "badge-under", "label": "🟠 Under target"},
    "blocked": {"cls": "status-blocked", "badge": "badge-blocked", "label": "🔴 Blocked"},
}


def render_manager(entries, notes):
    render_header(manager=True)
    st.markdown('<div class="manager-banner">Read-only overview of Saniya\'s daily progress. Private journal reflections are not shown here. You can leave an optional note on any specific day below.</div>', unsafe_allow_html=True)
    st.write("")

    gc1, gc2 = st.columns([0.5, 0.5])
    with gc1:
        with st.expander("📘 Guide — how to use this dashboard"):
            st.markdown(
                """
- **Weekly banners** group entries by week, with total hours logged that week.
- **Day cards** show each day's completed tasks, goals for the next day, and any notes/blockers Saniya flagged. The colored left edge and badge show status at a glance: 🟢 on track, 🟠 under target, 🔴 blocked.
- **Shared reflections** (key learnings / ideas) appear only when Saniya has chosen to make them visible here. Anything marked private in her journal stays private.
- **Leave a note on this day** — optional, per day. Expand it under any day card to add a comment or suggestion; it's saved with your name and timestamp and stays visible to both of you.
- **Downloads** below let you pull a full spreadsheet or a per-week Word report at any time — no need to ask Saniya to send one.
- This page is read-only for work entries — you can't edit Saniya's logged tasks or hours from here, only add notes.
                """
            )
    with gc2:
        with st.expander("⬇ Downloads"):
            render_downloads(entries, notes, key_prefix="mgr_")

    render_stats(entries)
    if not entries:
        st.info("Nothing logged yet.")
        return
    weeks = {}
    for iso in entries:
        weeks.setdefault(week_start(iso), []).append(iso)
    for ws in sorted(weeks.keys(), reverse=True):
        wdates = sorted(weeks[ws], reverse=True)
        total_hours = sum(float(entries[d].get("hours") or 0) for d in wdates)
        st.markdown(
            f'<div class="week-banner"><div><div class="wtitle">Week of {fmt_date(ws)}</div>'
            f'<div class="wsub">{len(wdates)} day{"s" if len(wdates)!=1 else ""} logged</div></div>'
            f'<div class="week-chip">{total_hours:g}h total</div></div>',
            unsafe_allow_html=True,
        )
        for iso in wdates:
            e = entries[iso]
            status = e.get("status", "on-track") if e.get("status") in STATUS_META else "on-track"
            meta = STATUS_META[status]
            st.markdown(f'<div class="day-card {meta["cls"]}">', unsafe_allow_html=True)
            st.markdown(
                f'<div class="day-head"><div><span class="day-date">{fmt_date(iso)}</span>'
                f'<span class="day-name">{day_name(iso)}</span></div>'
                f'<div><span class="badge {meta["badge"]}">{meta["label"]}</span>'
                f'<span class="hours-chip">{e.get("hours",0):g}h</span></div></div>',
                unsafe_allow_html=True,
            )
            st.markdown('<div class="section-label">Tasks completed</div>', unsafe_allow_html=True)
            render_task_list(e.get("tasks"))
            if e.get("goals"):
                st.markdown('<div class="section-label">Goals for next day</div>', unsafe_allow_html=True)
                st.write(e["goals"])
            if e.get("notes"):
                st.markdown('<div class="section-label">Notes / blockers</div>', unsafe_allow_html=True)
                st.write(e["notes"])
            if e.get("learning") or e.get("idea"):
                bits = []
                if e.get("learning"): bits.append(f"**Key learning:** {e['learning']}")
                if e.get("idea"): bits.append(f"**Idea:** {e['idea']}")
                st.markdown('<div class="section-label">Shared reflections</div>', unsafe_allow_html=True)
                for b in bits: st.write(b)

            day_notes = notes_for_day(notes, iso)
            for n in day_notes:
                st.markdown(
                    f'<div class="note-bubble"><div class="note-meta">{n.get("author","Manager")} · {n.get("ts","")}</div>{n.get("text","")}</div>',
                    unsafe_allow_html=True,
                )

            with st.expander("Leave a note on this day (optional)"):
                c1, c2 = st.columns([0.28, 0.72])
                author = c1.text_input("Name", value="Manager", key=f"author_{iso}", label_visibility="collapsed", placeholder="Your name")
                text = c2.text_input("Note", placeholder="Comment or suggestion — leave blank to skip", key=f"note_{iso}", label_visibility="collapsed")
                if st.button("Post note", key=f"post_{iso}"):
                    if text.strip():
                        add_manager_note(iso, author, text)
                        st.success("Note posted.")
                        st.rerun()
                    else:
                        st.caption("Nothing posted — note was left blank.")
            st.markdown('</div>', unsafe_allow_html=True)


def render_feedback(notes):
    st.header("Feedback")
    if not notes:
        st.info("No manager notes yet.")
        return
    for n in sorted(notes, key=lambda x:(x.get("date",""),x.get("ts","")), reverse=True):
        day = n.get("date")
        day_label = f"{fmt_date(day)} ({day_name(day)})" if day else "—"
        st.container(border=True).markdown(f"**{n.get('author','Manager')}** · {day_label} · {n.get('ts','')}\n\n{n.get('text','')}")


def build_word(entries, ws):
    from docx import Document
    from docx.shared import RGBColor
    doc = Document()
    title = doc.add_heading("Weekly Field Log", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1B,0x4F,0x91)
    doc.add_paragraph("Saniya · Intern · TraceLink").runs[0].italic = True
    doc.add_paragraph(f"Week of {fmt_date(ws)}").runs[0].bold = True
    dates = sorted([d for d in entries if week_start(d) == ws])
    doc.add_paragraph(f"Total hours logged: {sum(float(entries[d].get('hours') or 0) for d in dates):g}")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i,h in enumerate(["Date","Tasks completed","Goals for next day","Hrs","Status"]): table.rows[0].cells[i].text=h
    for d in dates:
        e=entries[d]; row=table.add_row().cells
        row[0].text=f"{fmt_date(d)} ({day_name(d)})"
        row[1].text=tasks_text(e.get('tasks')) or "—"
        row[2].text=e.get('goals') or "—"
        row[3].text=str(e.get('hours') or 0)
        row[4].text=(e.get('status') or 'on-track').replace('-',' ').title()
    notes=[n for n in load_notes() if n.get('date') in dates]
    if notes:
        doc.add_heading("Manager notes",level=2)
        for n in sorted(notes, key=lambda x: x.get('date','')):
            day = n.get('date')
            day_label = f"{fmt_date(day)}" if day else ""
            doc.add_paragraph(f"{day_label} — {n.get('author','Manager')} ({n.get('ts','')}): {n.get('text','')}")
    bio=BytesIO(); doc.save(bio); bio.seek(0); return bio


def build_xlsx(entries, notes):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = "Daily Log"
    headers = ["Date", "Day", "Tasks completed", "Goals for next day", "Hours", "Status", "Notes / blockers", "Key learning", "Idea", "Reflection"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1B4F91")
    for iso in sorted(entries):
        e = entries[iso]
        ws.append([iso, day_name(iso), tasks_text(e.get("tasks")), e.get("goals", ""), e.get("hours", 0), (e.get("status") or "on-track").replace("-", " ").title(), e.get("notes", ""), e.get("learning", ""), e.get("idea", ""), e.get("reflection", "")])
    widths = [14, 12, 55, 40, 10, 16, 40, 40, 40, 50]
    for i,w in enumerate(widths,1): ws.column_dimensions[chr(64+i)].width = w
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    wn = wb.create_sheet("Manager Notes")
    wn.append(["Date", "Author", "Timestamp", "Note"])
    for cell in wn[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0E7C86")
    for n in sorted(notes, key=lambda x: x.get("date", "")):
        wn.append([n.get("date", ""), n.get("author", "Manager"), n.get("ts", ""), n.get("text", "")])
    for col,w in zip("ABCD", [14,22,20,80]): wn.column_dimensions[col].width=w
    for row in wn.iter_rows():
        for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)
    bio = BytesIO(); wb.save(bio); bio.seek(0); return bio


def render_downloads(entries, notes, key_prefix=""):
    st.subheader("Downloads")
    xlsx = build_xlsx(entries, notes)
    st.download_button("⬇ Download full spreadsheet", data=xlsx, file_name="Field_Log_Full.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=f"{key_prefix}xlsx_full")
    weeks=sorted({week_start(d) for d in entries}, reverse=True)
    for ws in weeks:
        bio=build_word(entries,ws)
        st.download_button(f"⬇ Word report — week of {fmt_date(ws)}", data=bio, file_name=f"Field_Log_Week_{ws}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{key_prefix}word_{ws}")


def main():
    st.set_page_config(page_title="Field Log — Saniya", page_icon="FL", layout="wide", initial_sidebar_state="collapsed")
    inject_css()
    entries = load_entries()
    notes = load_notes()
    manager_mode = st.query_params.get("view", "") == "manager"
    if manager_mode:
        render_manager(entries, notes)
        return
    render_header(manager=False)
    render_stats(entries)
    with st.expander("Downloads"):
        render_downloads(entries, notes, key_prefix="own_")
    st.caption("Tip: share the Manager View link with your manager. Because this is hosted online, it works across different Wi-Fi/VPN networks.")
    tabs = st.tabs(["Today's Entry", "Ops Log", "Journal", "Manager View", "Feedback"])
    with tabs[0]: render_entry(entries)
    with tabs[1]: render_ops_log(entries)
    with tabs[2]: render_journal(entries)
    with tabs[3]:
        st.subheader("Manager View")
        st.markdown("Your manager can use the direct read-only URL: `?view=manager`")
        st.code(f"{st.context.headers.get('Host','YOUR-APP.streamlit.app')}/?view=manager", language="text")
        render_manager(entries, notes)
    with tabs[4]: render_feedback(notes)

if __name__ == "__main__":
    main()
