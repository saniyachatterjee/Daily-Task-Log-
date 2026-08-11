import streamlit as st
import pandas as pd
import json, os
from datetime import date, datetime, timedelta
from io import BytesIO

# ── TraceLink Anthem brand palette (opus.tracelink.com/.../anthem) ─────────
EVERGREEN   = "#006341"   # primary
DEEP_OCEAN  = "#003D4C"   # secondary / backdrop
CYAN        = "#00A7B5"   # primary CTA
CYAN_SOFT   = "#E5F6F8"
COBALT      = "#005587"   # secondary CTA
COBALT_SOFT = "#E5EDF3"
EMERALD     = "#008755"   # success / on-track
EMERALD_SOFT= "#E5F3ED"
CITRON      = "#ECE81A"   # decorative accent
RED         = "#be0000"   # failure / blocked
RED_SOFT    = "#f9e5e5"
ORANGE      = "#ff6e3a"   # warning / under target
ORANGE_SOFT = "#fff0eb"
GRAY_900    = "#252525"
GRAY_700    = "#505050"
GRAY_500    = "#7C7C7C"
GRAY_300    = "#BDBDBD"
GRAY_100    = "#E8E8E8"
GRAY_50     = "#F4F4F4"

STATUS_META = {
    "on-track": {"label": "On track",     "color": EMERALD, "soft": EMERALD_SOFT, "emoji": "🟢"},
    "under":    {"label": "Under target",  "color": ORANGE,  "soft": ORANGE_SOFT,  "emoji": "🟠"},
    "blocked":  {"label": "Blocked",       "color": RED,     "soft": RED_SOFT,     "emoji": "🔴"},
}

DATA_DIR = "data"
ENTRIES_FILE = os.path.join(DATA_DIR, "entries.json")
NOTES_FILE = os.path.join(DATA_DIR, "manager_notes.json")
os.makedirs(DATA_DIR, exist_ok=True)

SEED = {
    "2026-08-03": {"date":"2026-08-03","tasks":"- Reviewed, edited and sent 13 T18 documents to Alok\n- Worked on Conval and Excellis Audit Report spreadsheet for Pooja","goals":"- Work on more T18 documents","hours":7,"status":"on-track","notes":"","learning":"","idea":"","reflection":"","private":True,"managerNote":""},
    "2026-08-04": {"date":"2026-08-04","tasks":"- Reviewed, edited and sent all remaining T18 documents\n- Finished some Litmos courses\n- Meeting with Alok, Pooja, Deepak on Partner Scorecard Dashboard","goals":"Set clearer goal for next day","hours":6,"status":"on-track","notes":"Which Litmos course, how much time per activity — need to specify daily goals better","learning":"","idea":"","reflection":"","private":True,"managerNote":""},
    "2026-08-05": {"date":"2026-08-05","tasks":"Litmos courses completed: FOSCM, QMS Privacy & Data Protection, NS Overview for New Hires, QMS CS01 Technical Support, Preventing Harassment (Global Workplace), QMS Data Integrity Policy","goals":"- Complete QMS PS01 course\n- Complete more Litmos courses\n- Prepare and attend SPE Monthly Meet\n- Attend Use Org (Services Training) Meet","hours":7,"status":"on-track","notes":"","learning":"","idea":"","reflection":"","private":True,"managerNote":""},
    "2026-08-06": {"date":"2026-08-06","tasks":"Meetings attended: SPE Monthly Meeting, Use Org (Services Training) Meet.\nLitmos completed: QMS PS01, SVC Use Org Training, Intro to GS1 Standards, OPUS No Code Fundamentals, MINT External Manufacturing Configuration","goals":"- Complete all remaining Litmos courses before 8 Aug deadline\n- Weekly catchup with manager","hours":7,"status":"under","notes":"Litmos deadline is tight — running behind on hours logged mid-week","learning":"","idea":"","reflection":"","private":True,"managerNote":""},
}

# ── Storage helpers ──────────────────────────────────────────────────────
def load_entries():
    if not os.path.exists(ENTRIES_FILE):
        with open(ENTRIES_FILE, "w") as f:
            json.dump(SEED, f, indent=2)
        return dict(SEED)
    with open(ENTRIES_FILE) as f:
        return json.load(f)

def save_entries(entries):
    with open(ENTRIES_FILE, "w") as f:
        json.dump(entries, f, indent=2)

def load_notes():
    if not os.path.exists(NOTES_FILE):
        return []
    with open(NOTES_FILE) as f:
        return json.load(f)

def save_notes(notes):
    with open(NOTES_FILE, "w") as f:
        json.dump(notes, f, indent=2)

def week_start(d: date) -> date:
    return d - timedelta(days=d.weekday())

def day_name(iso): return datetime.strptime(iso, "%Y-%m-%d").strftime("%a")
def fmt_date(iso): return datetime.strptime(iso, "%Y-%m-%d").strftime("%d %b %Y")

# ── Page setup + TraceLink-styled CSS ───────────────────────────────────
st.set_page_config(page_title="Field Log — Saniya", page_icon="🛰️", layout="wide")

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:wght@600;700&family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {{ font-family:'Inter',sans-serif; }}

.stApp {{
  background: linear-gradient(135deg, {EVERGREEN} 0%, {DEEP_OCEAN} 100%);
}}
section[data-testid="stSidebar"] {{
  background: {DEEP_OCEAN};
}}
section[data-testid="stSidebar"] * {{ color: #EAF3F1 !important; }}

#MainMenu, footer, header {{ visibility: hidden; }}

.tl-banner{{
  background: rgba(255,255,255,0.06);
  border: 1px solid rgba(255,255,255,0.14);
  border-radius: 16px;
  padding: 22px 28px;
  margin-bottom: 22px;
  backdrop-filter: blur(6px);
}}
.tl-eyebrow{{
  font-size:11.5px; letter-spacing:.12em; text-transform:uppercase; font-weight:700;
  color:{CITRON}; display:flex; align-items:center; gap:8px;
}}
.tl-eyebrow .dot{{width:7px;height:7px;border-radius:50%;background:#4ADE80;box-shadow:0 0 0 4px rgba(74,222,128,.25);}}
.tl-title{{ font-family:'Fraunces',serif; font-size:36px; font-weight:700; color:#fff; margin:6px 0 2px; }}
.tl-sub{{ color:#CFE3DF; font-size:14px; }}

.tl-card{{
  background:#fff; border-radius:16px; padding:22px 26px; margin-bottom:16px;
  box-shadow:0 8px 24px rgba(0,0,0,0.18);
}}

div[data-testid="stMetric"]{{
  background:#fff; border-radius:14px; padding:12px 16px; box-shadow:0 6px 16px rgba(0,0,0,.15);
}}
div[data-testid="stMetricValue"]{{ color:{EVERGREEN}; font-weight:700; }}

.stTabs [data-baseweb="tab-list"]{{ gap:4px; }}
.stTabs [data-baseweb="tab"]{{
  background:rgba(255,255,255,0.08); border-radius:10px 10px 0 0; color:#EAF3F1; font-weight:600;
  padding:10px 18px;
}}
.stTabs [aria-selected="true"]{{
  background:#fff !important; color:{EVERGREEN} !important;
}}

.stButton>button{{
  background:{CYAN}; color:#fff; border:none; border-radius:9px; font-weight:700; padding:0.55em 1.3em;
  transition: transform .1s, box-shadow .15s;
}}
.stButton>button:hover{{ box-shadow:0 6px 16px rgba(0,167,181,.4); transform:translateY(-1px); }}
.stDownloadButton>button{{
  background:{COBALT}; color:#fff; border:none; border-radius:9px; font-weight:700;
}}

.status-badge{{
  display:inline-flex; align-items:center; gap:6px; padding:4px 12px; border-radius:20px;
  font-weight:700; font-size:12.5px;
}}

.tl-node-track{{ display:flex; align-items:center; overflow-x:auto; padding:6px 2px; }}
.tl-node{{ display:flex; flex-direction:column; align-items:center; margin-right:6px; flex-shrink:0; }}
.tl-node .dot{{ width:20px; height:20px; border-radius:50%; border:2.5px solid #fff; box-shadow:0 2px 5px rgba(0,0,0,.25); }}
.tl-node .lbl{{ font-size:9.5px; color:#EAF3F1; margin-top:4px; font-weight:600; }}
.tl-line{{ width:22px; height:2px; background:rgba(255,255,255,0.35); margin:0 -2px 16px; flex-shrink:0; }}

.journal-card{{
  background:#fff; border-left:5px solid {COBALT}; border-radius:12px; padding:16px 20px; margin-bottom:12px;
  box-shadow:0 4px 14px rgba(0,0,0,.12);
}}
.journal-card .jlbl{{ font-size:11px; font-weight:700; color:{COBALT}; text-transform:uppercase; letter-spacing:.05em; margin-top:8px;}}
.journal-card .jlbl:first-child{{margin-top:0;}}

.mgr-note-existing{{
  background:{RED_SOFT}; border:1.5px solid #f3c8c8; border-radius:10px; padding:12px 14px; margin-top:10px; color:#7c1f1f; font-size:14px;
}}

/* ── Seamless glass panels for bordered containers (forms, comment boxes) ── */
div[data-testid="stVerticalBlockBorderWrapper"]{{
  background: rgba(255,255,255,0.07);
  border: 1px solid rgba(255,255,255,0.16) !important;
  border-radius: 16px !important;
  backdrop-filter: blur(6px);
  box-shadow: 0 8px 24px rgba(0,0,0,0.15);
}}
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] p,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] li,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] strong,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h1,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h2,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h3,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h4,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h5,
div[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stWidgetLabel"] p,
div[data-testid="stVerticalBlockBorderWrapper"] .stRadio label p,
div[data-testid="stVerticalBlockBorderWrapper"] .stCheckbox label p,
div[data-testid="stVerticalBlockBorderWrapper"] .stCaptionContainer{{
  color:#F4F4F4 !important;
}}

/* ── Force legible white text inside every input / textarea / dropdown ── */
.stTextArea textarea, .stTextInput input, .stNumberInput input, .stDateInput input{{
  background: rgba(255,255,255,0.09) !important;
  color:#FFFFFF !important;
  border:1px solid rgba(255,255,255,0.28) !important;
  border-radius:10px !important;
  caret-color:#FFFFFF;
}}
.stTextArea textarea::placeholder, .stTextInput input::placeholder{{
  color: rgba(255,255,255,0.55) !important;
}}
div[data-baseweb="select"] > div{{
  background: rgba(255,255,255,0.09) !important;
  border:1px solid rgba(255,255,255,0.28) !important;
  border-radius:10px !important;
}}
div[data-baseweb="select"] *{{ color:#FFFFFF !important; }}
div[data-baseweb="popover"] li{{ color:{GRAY_900} !important; }}

/* ── Plain headings/labels sitting directly on the gradient ── */
.tl-plain-heading{{ color:#F4F4F4; font-weight:700; margin-bottom:2px; }}

div[data-testid="stMetric"] label{{ color:{GRAY_700} !important; }}

.task-bullet{{
  background: rgba(255,255,255,0.08);
  border: 1px solid rgba(255,255,255,0.16);
  border-radius: 8px;
  padding: 8px 12px;
  margin-bottom: 6px;
  color: #FFFFFF;
  font-size: 14.5px;
}}
</style>
""", unsafe_allow_html=True)

entries = load_entries()
notes = load_notes()

# ── Sidebar navigation ───────────────────────────────────────────────────
st.sidebar.markdown("### 🛰️ Field Log")
st.sidebar.caption("Network Success Operations · TraceLink")
view = st.sidebar.radio("View", ["📝 My Log", "👩‍💼 Manager View"], label_visibility="collapsed")
st.sidebar.divider()
st.sidebar.caption("Built for Saniya · daily ops + working journal")

# ── Shared stats ─────────────────────────────────────────────────────────
def compute_stats():
    today = date.today()
    ws = week_start(today)
    we = ws + timedelta(days=4)
    week_hours = on_t = under = blocked = week_count = 0
    for iso, e in entries.items():
        d = datetime.strptime(iso, "%Y-%m-%d").date()
        if ws <= d <= we:
            week_hours += float(e.get("hours") or 0)
            week_count += 1
            s = e.get("status", "on-track")
            if s == "on-track": on_t += 1
            elif s == "under": under += 1
            elif s == "blocked": blocked += 1
    dates_with_data = sorted([d for d in entries if entries[d].get("tasks") or entries[d].get("hours")])
    streak = 0
    if dates_with_data:
        streak = 1
        for i in range(len(dates_with_data)-1, 0, -1):
            cur = datetime.strptime(dates_with_data[i], "%Y-%m-%d").date()
            prev = datetime.strptime(dates_with_data[i-1], "%Y-%m-%d").date()
            if (cur - prev).days <= 3:
                streak += 1
            else:
                break
    return streak, week_hours, week_count, on_t, under, blocked

def render_banner(subtitle):
    st.markdown(f"""
    <div class="tl-banner">
      <div class="tl-eyebrow"><span class="dot"></span>Network Success Operations · TraceLink</div>
      <div class="tl-title">Field Log</div>
      <div class="tl-sub">{subtitle}</div>
    </div>
    """, unsafe_allow_html=True)

def render_stat_row():
    streak, wh, wc, ot, un, bl = compute_stats()
    cols = st.columns(6)
    cols[0].metric("Day streak", streak)
    cols[1].metric("Hrs this week", wh)
    cols[2].metric("Entries this week", wc)
    cols[3].metric("On track", ot)
    cols[4].metric("Under target", un)
    cols[5].metric("Blocked", bl)

def render_signal_strip():
    dates_sorted = sorted(entries.keys())
    start_d = datetime.strptime(dates_sorted[0], "%Y-%m-%d").date() if dates_sorted else date.today()
    end_d = date.today() + timedelta(days=4)
    nodes = []
    cur = start_d
    while cur <= end_d:
        if cur.weekday() < 5:
            nodes.append(cur.isoformat())
        cur += timedelta(days=1)
    today_iso = date.today().isoformat()
    html = '<div class="tl-node-track">'
    for i, iso in enumerate(nodes):
        e = entries.get(iso)
        status = e.get("status", "on-track") if e else None
        color = STATUS_META[status]["color"] if status else "rgba(255,255,255,0.25)"
        ring = f"box-shadow:0 0 0 3px {CYAN_SOFT}66;" if iso == today_iso else ""
        short = datetime.strptime(iso, "%Y-%m-%d").strftime("%d %b")
        html += f'<div class="tl-node"><div class="dot" style="background:{color};{ring}"></div><div class="lbl">{short}</div></div>'
        if i < len(nodes) - 1:
            html += '<div class="tl-line"></div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)

def xlsx_bytes():
    rows = []
    for iso in sorted(entries.keys()):
        e = entries[iso]
        rows.append({
            "Date": iso, "Day": day_name(iso),
            "Tasks completed": e.get("tasks",""), "Goals for tomorrow": e.get("goals",""),
            "Hours spent": e.get("hours",0), "Notes/blockers": e.get("notes",""),
            "Status": STATUS_META.get(e.get("status","on-track"),{}).get("label",""),
            "Key learning": e.get("learning",""), "Idea/spark": e.get("idea",""),
            "Reflection": e.get("reflection",""), "Manager note": e.get("managerNote",""),
        })
    df = pd.DataFrame(rows)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Field Log")
        ws = writer.sheets["Field Log"]
        for i, col in enumerate(df.columns, 1):
            ws.column_dimensions[chr(64+i) if i <= 26 else "A"].width = max(14, min(40, df[col].astype(str).map(len).max() if len(df) else 14))
        ws.freeze_panes = "A2"
    buf.seek(0)
    return buf

def docx_bytes_week(ws_iso):
    from docx import Document
    from docx.shared import RGBColor
    week_dates = sorted([d for d in entries if week_start(datetime.strptime(d,"%Y-%m-%d").date()).isoformat() == ws_iso])
    doc = Document()
    h = doc.add_heading("Weekly Field Log", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x00, 0x63, 0x41)
    doc.add_paragraph("Saniya · Network Success Operations · TraceLink").runs[0].italic = True
    doc.add_paragraph(f"Week of {fmt_date(ws_iso)}").runs[0].bold = True
    total = sum(float(entries[d].get("hours") or 0) for d in week_dates)
    doc.add_paragraph(f"Total hours logged: {total}")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, hd in enumerate(["Date","Tasks completed","Goals for next day","Hrs","Status"]):
        table.rows[0].cells[i].text = hd
    for d in week_dates:
        e = entries[d]
        row = table.add_row().cells
        row[0].text = f"{fmt_date(d)} ({day_name(d)})"
        row[1].text = e.get("tasks","") or "—"
        row[2].text = e.get("goals","") or "—"
        row[3].text = str(e.get("hours","") or "0")
        row[4].text = STATUS_META.get(e.get("status","on-track"),{}).get("label","")
    learn = [(d, entries[d]) for d in week_dates if entries[d].get("learning") or entries[d].get("idea")]
    if learn:
        doc.add_heading("Key learnings & ideas", level=2)
        for d, e in learn:
            if e.get("learning"): doc.add_paragraph(f"{fmt_date(d)} — Learning: {e['learning']}", style="List Bullet")
            if e.get("idea"): doc.add_paragraph(f"{fmt_date(d)} — Idea: {e['idea']}", style="List Bullet")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════
# MY LOG VIEW
# ══════════════════════════════════════════════════════════════════════
if view == "📝 My Log":
    render_banner("Daily ops tracking & working journal")
    render_stat_row()
    st.write("")
    st.markdown('<div class="tl-plain-heading" style="font-size:15px;">Signal strip</div>', unsafe_allow_html=True)
    render_signal_strip()
    st.write("")

    tab1, tab2, tab3 = st.tabs(["Today's Entry", "Ops Log", "Journal"])

    with tab1:
        with st.container(border=True):
            sel_date = st.date_input("Date", value=date.today())
            iso = sel_date.isoformat()
            e = entries.get(iso, {})

            if e.get("managerNote"):
                st.markdown(f'<div class="mgr-note-existing">💬 <b>Manager\'s note on this day:</b><br>{e["managerNote"]}</div>', unsafe_allow_html=True)
                st.write("")

            st.markdown(f"**{day_name(iso)}, {fmt_date(iso)}**")
            st.markdown("##### Ops — visible to manager")

            # ── Tasks completed — bulleted add/remove list ──────────────
            task_list_key = f"task_list_{iso}"
            task_input_key = f"task_input_{iso}"
            if task_list_key not in st.session_state:
                st.session_state[task_list_key] = [
                    ln.strip().lstrip("-•").strip()
                    for ln in e.get("tasks", "").split("\n") if ln.strip()
                ]

            def _add_task(list_key=task_list_key, input_key=task_input_key):
                val = st.session_state.get(input_key, "").strip()
                if val:
                    st.session_state[list_key].append(val)
                    st.session_state[input_key] = ""

            def _remove_task(idx, list_key=task_list_key):
                if 0 <= idx < len(st.session_state[list_key]):
                    st.session_state[list_key].pop(idx)

            st.markdown("**Tasks completed**")
            if st.session_state[task_list_key]:
                for i, t in enumerate(st.session_state[task_list_key]):
                    tcol1, tcol2 = st.columns([11, 1])
                    tcol1.markdown(f'<div class="task-bullet">• {t}</div>', unsafe_allow_html=True)
                    tcol2.button("✕", key=f"rm_{iso}_{i}", on_click=_remove_task, args=(i,))
            else:
                st.caption("No tasks added yet — add your first one below.")
            ac1, ac2 = st.columns([11, 1])
            ac1.text_input("Add a task", key=task_input_key, on_change=_add_task,
                            placeholder="e.g. Reviewed and sent 6 T18 documents", label_visibility="collapsed")
            ac2.button("➕", key=f"add_{iso}", on_click=_add_task)
            tasks = "\n".join(f"- {t}" for t in st.session_state[task_list_key])

            goals = st.text_area("Goals for tomorrow", value=e.get("goals",""), placeholder="- Finish remaining Litmos courses")
            c1, c2 = st.columns(2)
            with c1:
                hours = st.number_input("Hours spent", min_value=0.0, max_value=16.0, step=0.5, value=float(e.get("hours") or 0))
            with c2:
                status_labels = {v["emoji"]+" "+v["label"]: k for k, v in STATUS_META.items()}
                cur_status = e.get("status","on-track")
                cur_label = STATUS_META[cur_status]["emoji"]+" "+STATUS_META[cur_status]["label"]
                chosen = st.radio("Status", list(status_labels.keys()), index=list(status_labels.keys()).index(cur_label), horizontal=True)
                status = status_labels[chosen]
            notes_val = st.text_area("Notes / blockers (optional)", value=e.get("notes",""))

            st.markdown("##### Journal — your working reflection")
            learning = st.text_area("Key learning today", value=e.get("learning",""))
            idea = st.text_area("Idea / spark (optional)", value=e.get("idea",""))
            reflection = st.text_area("Thoughts / reflection", value=e.get("reflection",""))
            private = st.checkbox("Keep this reflection private (excluded from Manager View & exports' comment)", value=e.get("private", True))

            if st.button("💾 Save entry", type="primary"):
                entries[iso] = {
                    "date": iso, "tasks": tasks, "goals": goals, "hours": hours,
                    "status": status, "notes": notes_val, "learning": learning,
                    "idea": idea, "reflection": reflection, "private": private,
                    "managerNote": e.get("managerNote","")
                }
                save_entries(entries)
                st.success("Saved ✓")
                st.rerun()

    with tab2:
        if entries:
            df = pd.DataFrame([{
                "Date": fmt_date(d), "Day": day_name(d), "Tasks": entries[d].get("tasks","")[:80],
                "Goals": entries[d].get("goals","")[:80], "Hrs": entries[d].get("hours",0),
                "Status": STATUS_META.get(entries[d].get("status","on-track"),{}).get("label",""),
                "Notes": entries[d].get("notes","")[:60],
            } for d in sorted(entries.keys(), reverse=True)])
            st.dataframe(df, use_container_width=True, hide_index=True)
        else:
            st.info("No entries yet.")

    with tab3:
        j_dates = [d for d in sorted(entries.keys(), reverse=True) if entries[d].get("learning") or entries[d].get("idea") or entries[d].get("reflection")]
        if not j_dates:
            st.info("No journal reflections yet.")
        for d in j_dates:
            e = entries[d]
            tag = ' <span style="float:right;font-size:10px;background:#E8E8E8;color:#505050;padding:2px 8px;border-radius:8px;">PRIVATE</span>' if e.get("reflection") and e.get("private", True) else ""
            html = f'<div class="journal-card"><div style="font-size:12px;color:{GRAY_500};font-weight:600;">{fmt_date(d)} · {day_name(d)}{tag}</div>'
            if e.get("learning"): html += f'<div class="jlbl">Key learning</div><div>{e["learning"]}</div>'
            if e.get("idea"): html += f'<div class="jlbl">Idea / spark</div><div>{e["idea"]}</div>'
            if e.get("reflection"): html += f'<div class="jlbl">Reflection</div><div>{e["reflection"]}</div>'
            html += '</div>'
            st.markdown(html, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════
# MANAGER VIEW
# ══════════════════════════════════════════════════════════════════════
else:
    render_banner("Manager overview — read-only, with room for your feedback")
    render_stat_row()
    st.write("")
    st.markdown('<div class="tl-plain-heading" style="font-size:15px;">Signal strip</div>', unsafe_allow_html=True)
    render_signal_strip()
    st.write("")

    st.markdown('<div class="tl-card">📌 <b>Private reflections are hidden here.</b> Use the exports below for a Word weekly report or full spreadsheet.</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        st.download_button("⬇ Download full spreadsheet (.xlsx)", data=xlsx_bytes(),
            file_name="Field_Log_Saniya.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with c2:
        weeks_available = sorted({week_start(datetime.strptime(d,"%Y-%m-%d").date()).isoformat() for d in entries}, reverse=True)
        if weeks_available:
            wk_choice = st.selectbox("Week for Word report", weeks_available, format_func=fmt_date)
            st.download_button("⬇ Download Word report for this week", data=docx_bytes_week(wk_choice),
                file_name=f"Field_Log_Week_{wk_choice}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.write("")
    st.markdown('<div class="tl-plain-heading" style="font-size:24px;margin-top:6px;">Weekly overview</div>', unsafe_allow_html=True)

    weeks = {}
    for d in entries:
        ws = week_start(datetime.strptime(d, "%Y-%m-%d").date()).isoformat()
        weeks.setdefault(ws, []).append(d)

    for ws in sorted(weeks.keys(), reverse=True):
        wdates = sorted(weeks[ws])
        total = sum(float(entries[d].get("hours") or 0) for d in wdates)
        with st.container(border=True):
            st.markdown(f"**Week of {fmt_date(ws)}** — {total}h logged")
            rows = []
            for d in wdates:
                e = entries[d]
                extra = ""
                if e.get("learning"): extra += f"Learning: {e['learning']}  "
                if e.get("idea"): extra += f"Idea: {e['idea']}"
                rows.append({
                    "Date": f"{fmt_date(d)} ({day_name(d)})", "Tasks": e.get("tasks",""),
                    "Goals": e.get("goals",""), "Hrs": e.get("hours",0),
                    "Status": STATUS_META.get(e.get("status","on-track"),{}).get("label",""),
                    "Learnings / ideas": extra,
                })
            st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

    st.write("")
    st.markdown('<div class="tl-plain-heading" style="font-size:24px;margin-top:6px;">Leave a comment on a specific day</div>', unsafe_allow_html=True)
    with st.container(border=True):
        target_date = st.selectbox("Day", sorted(entries.keys(), reverse=True), format_func=lambda d: f"{fmt_date(d)} ({day_name(d)})")
        existing = entries.get(target_date, {}).get("managerNote", "")
        comment = st.text_area("Your thoughts, ideas, or suggestions for this day", value=existing)
        if st.button("💬 Save comment"):
            entries[target_date]["managerNote"] = comment
            save_entries(entries)
            st.success("Comment saved — Saniya will see it on this day's entry.")
            st.rerun()

    st.write("")
    st.markdown('<div class="tl-plain-heading" style="font-size:24px;margin-top:6px;">💡 Add a thought or suggestion</div>', unsafe_allow_html=True)
    with st.container(border=True):
        new_note = st.text_area("Not tied to a specific day — general feedback for Saniya", key="general_note",
                                 placeholder="e.g. Great progress on the Litmos deadline this week — let's discuss the T18 backlog on Friday.")
        if st.button("➕ Add"):
            if new_note.strip():
                notes.append({"text": new_note.strip(), "ts": datetime.now().strftime("%d %b %Y, %H:%M")})
                save_notes(notes)
                st.success("Added ✓")
                st.rerun()
        if notes:
            for n in reversed(notes[-10:]):
                st.markdown(f'<div class="journal-card"><div style="font-size:11px;color:{GRAY_500};">{n["ts"]}</div>{n["text"]}</div>', unsafe_allow_html=True)
